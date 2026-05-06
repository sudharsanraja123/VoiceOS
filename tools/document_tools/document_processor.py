"""
Document Processor - Safe wrapper for Agent Zero document processing
Maintains VoiceOS security boundaries while leveraging imported capabilities
"""

import os
import logging
from typing import Optional, Dict, Any, List
from datetime import datetime
from pathlib import Path

# VoiceOS Tools - Native implementation
import os
import logging
from typing import Optional, Dict, Any, List
from datetime import datetime
from pathlib import Path

import PyPDF2
try:
    import docx
except ImportError:
    docx = None
import chardet

from core.config import config
from permissions.permission_engine import PermissionLevel, check_permission


class DocumentProcessor:
    """
    Safe wrapper for document processing with validation and sandboxing
    """
    
    def __init__(self, workspace_root: Optional[str] = None):
        self.workspace_root = Path(workspace_root) if workspace_root else config.project_root / "workspace"
        self.workspace_root.mkdir(exist_ok=True)
        self.logger = logging.getLogger(__name__)
        
        # Security constraints
        self.allowed_extensions = ['.txt', '.md', '.pdf', '.docx', '.doc', '.rtf']
        self.max_file_size_mb = 50
        self.max_content_length = 1024 * 1024  # 1MB for processing
        
    def _validate_file(self, file_path: str) -> Path:
        """Validate file path and properties"""
        try:
            resolved_path = Path(file_path).resolve()
            
            # Ensure path is within workspace
            if not str(resolved_path).startswith(str(self.workspace_root.resolve())):
                raise PermissionError(f"File {file_path} is outside workspace bounds")
            
            # Check if file exists
            if not resolved_path.exists():
                raise FileNotFoundError(f"File not found: {file_path}")
            
            # Check file extension
            if resolved_path.suffix.lower() not in self.allowed_extensions:
                raise ValueError(f"File type {resolved_path.suffix} not allowed")
            
            # Check file size
            file_size_mb = resolved_path.stat().st_size / (1024 * 1024)
            if file_size_mb > self.max_file_size_mb:
                raise ValueError(f"File too large: {file_size_mb:.2f}MB (max: {self.max_file_size_mb}MB)")
            
            return resolved_path
            
        except Exception as e:
            self.logger.error(f"File validation failed for {file_path}: {e}")
            raise ValueError(f"Invalid file: {e}")
    
    def _log_operation(self, operation: str, file_path: str, result: Any, error: Optional[str] = None):
        """Log all document operations"""
        log_entry = {
            "timestamp": datetime.now().isoformat(),
            "operation": operation,
            "file_path": file_path,
            "result": str(result)[:200],  # Truncate long results
            "error": error
        }
        
        log_file = self.workspace_root / "logs" / "document_operations.log"
        log_file.parent.mkdir(exist_ok=True)
        
        with open(log_file, "a") as f:
            f.write(f"{log_entry}\n")
    
    @check_permission(PermissionLevel.LOW)
    def extract_text(self, file_path: str) -> Dict[str, Any]:
        """Extract text from document"""
        try:
            validated_path = self._validate_file(file_path)
            
            # VoiceOS native document extraction
            text_content = ""
            
            if validated_path.suffix.lower() == '.txt':
                with open(validated_path, 'r', encoding='utf-8') as f:
                    text_content = f.read()
            elif validated_path.suffix.lower() == '.pdf':
                try:
                    with open(validated_path, 'rb') as f:
                        pdf_reader = PyPDF2.PdfReader(f)
                        for page in pdf_reader.pages:
                            text_content += page.extract_text() + "\n"
                except Exception as e:
                    raise ValueError(f"PDF extraction failed: {e}")
            elif validated_path.suffix.lower() == '.docx':
                if docx is None:
                    raise ValueError("python-docx package not installed. Cannot process DOCX files.")
                try:
                    doc = docx.Document(str(validated_path))
                    for paragraph in doc.paragraphs:
                        text_content += paragraph.text + "\n"
                except Exception as e:
                    raise ValueError(f"DOCX extraction failed: {e}")
            else:
                # Try to read as text file with encoding detection
                with open(validated_path, 'rb') as f:
                    raw_data = f.read()
                    detected = chardet.detect(raw_data)
                    encoding = detected.get('encoding', 'utf-8')
                    text_content = raw_data.decode(encoding, errors='ignore')
            
            # Truncate if too long
            if len(text_content) > self.max_content_length:
                text_content = text_content[:self.max_content_length] + "\n...[truncated]"
            
            result = {
                "text": text_content,
                "file_name": validated_path.name,
                "file_size": validated_path.stat().st_size,
                "file_type": validated_path.suffix
            }
            
            self._log_operation("extract_text", file_path, "success")
            return result
            
        except Exception as e:
            self._log_operation("extract_text", file_path, "failed", str(e))
            raise
    
    @check_permission(PermissionLevel.LOW)
    def summarize_document(self, file_path: str, max_length: int = 500) -> Dict[str, Any]:
        """Generate document summary"""
        try:
            validated_path = self._validate_file(file_path)
            
            # Extract text first
            text_content = ""
            
            if validated_path.suffix.lower() == '.txt':
                with open(validated_path, 'r', encoding='utf-8') as f:
                    text_content = f.read()
            elif validated_path.suffix.lower() == '.pdf':
                try:
                    with open(validated_path, 'rb') as f:
                        pdf_reader = PyPDF2.PdfReader(f)
                        for page in pdf_reader.pages:
                            text_content += page.extract_text() + "\n"
                except Exception as e:
                    raise ValueError(f"PDF extraction failed: {e}")
            elif validated_path.suffix.lower() == '.docx':
                if docx is None:
                    raise ValueError("python-docx package not installed. Cannot process DOCX files.")
                try:
                    doc = docx.Document(str(validated_path))
                    for paragraph in doc.paragraphs:
                        text_content += paragraph.text + "\n"
                except Exception as e:
                    raise ValueError(f"DOCX extraction failed: {e}")
            
            # VoiceOS native summarization
            # Simple extractive summarization
            sentences = text_content.split('. ')
            if len(sentences) > max_length // 100:  # Rough estimate
                summary_sentences = sentences[:max_length // 100]
                summary = '. '.join(summary_sentences)
            else:
                summary = text_content[:max_length]
            
            result = {
                "summary": summary,
                "file_name": validated_path.name,
                "original_length": len(text_content),
                "summary_length": len(summary)
            }
            
            self._log_operation("summarize_document", file_path, "success")
            return result
            
        except Exception as e:
            self._log_operation("summarize_document", file_path, "failed", str(e))
            raise
    
    @check_permission(PermissionLevel.LOW)
    def search_in_document(self, file_path: str, query: str) -> Dict[str, Any]:
        """Search for text within document"""
        try:
            validated_path = self._validate_file(file_path)
            
            # Validate query
            if not query or len(query.strip()) == 0:
                raise ValueError("Search query cannot be empty")
            
            if len(query) > 200:
                raise ValueError("Search query too long")
            
            # Extract text first
            text_content = ""
            
            if validated_path.suffix.lower() == '.txt':
                with open(validated_path, 'r', encoding='utf-8') as f:
                    text_content = f.read()
            elif validated_path.suffix.lower() == '.pdf':
                try:
                    with open(validated_path, 'rb') as f:
                        pdf_reader = PyPDF2.PdfReader(f)
                        for page in pdf_reader.pages:
                            text_content += page.extract_text() + "\n"
                except Exception as e:
                    raise ValueError(f"PDF extraction failed: {e}")
            elif validated_path.suffix.lower() == '.docx':
                if docx is None:
                    raise ValueError("python-docx package not installed. Cannot process DOCX files.")
                try:
                    doc = docx.Document(str(validated_path))
                    for paragraph in doc.paragraphs:
                        text_content += paragraph.text + "\n"
                except Exception as e:
                    raise ValueError(f"DOCX extraction failed: {e}")
            
            # VoiceOS native text search
            matches = []
            lines = text_content.split('\n')
            for i, line in enumerate(lines):
                if query.lower() in line.lower():
                    matches.append({
                        "line_number": i + 1,
                        "line_content": line.strip(),
                        "match_position": line.lower().find(query.lower())
                    })
            
            search_results = {
                "matches": matches[:50],  # Limit results
                "total_matches": len(matches)
            }
            
            result = {
                "query": query,
                "matches": search_results.get("matches", []),
                "total_matches": search_results.get("total_matches", 0),
                "file_name": validated_path.name
            }
            
            self._log_operation("search_in_document", file_path, f"found {result['total_matches']} matches")
            return result
            
        except Exception as e:
            self._log_operation("search_in_document", file_path, "failed", str(e))
            raise
    
    @check_permission(PermissionLevel.MEDIUM)
    def analyze_document(self, file_path: str) -> Dict[str, Any]:
        """Analyze document structure and metadata"""
        try:
            validated_path = self._validate_file(file_path)
            
            # VoiceOS native document analysis
            file_stats = validated_path.stat()
            
            # Extract basic text for analysis
            text_content = ""
            
            if validated_path.suffix.lower() == '.txt':
                with open(validated_path, 'r', encoding='utf-8') as f:
                    text_content = f.read()
            elif validated_path.suffix.lower() == '.pdf':
                try:
                    with open(validated_path, 'rb') as f:
                        pdf_reader = PyPDF2.PdfReader(f)
                        for page in pdf_reader.pages:
                            text_content += page.extract_text() + "\n"
                except Exception as e:
                    text_content = f"PDF extraction failed: {e}"
            elif validated_path.suffix.lower() == '.docx':
                try:
                    doc = docx.Document(str(validated_path))
                    for paragraph in doc.paragraphs:
                        text_content += paragraph.text + "\n"
                except Exception as e:
                    text_content = f"DOCX extraction failed: {e}"
            
            # Basic analysis
            word_count = len(text_content.split())
            char_count = len(text_content)
            line_count = len(text_content.split('\n'))
            
            analysis = {
                "structure": {
                    "file_type": validated_path.suffix,
                    "encoding": "utf-8" if validated_path.suffix == '.txt' else "binary"
                },
                "metadata": {
                    "file_size": file_stats.st_size,
                    "created_time": file_stats.st_ctime,
                    "modified_time": file_stats.st_mtime
                },
                "statistics": {
                    "word_count": word_count,
                    "character_count": char_count,
                    "line_count": line_count
                }
            }
            
            # Sanitize analysis results
            sanitized_analysis = {
                "file_name": validated_path.name,
                "file_type": validated_path.suffix,
                "file_size": validated_path.stat().st_size,
                "structure": analysis.get("structure", {}),
                "metadata": analysis.get("metadata", {}),
                "statistics": analysis.get("statistics", {})
            }
            
            self._log_operation("analyze_document", file_path, "success")
            return sanitized_analysis
            
        except Exception as e:
            self._log_operation("analyze_document", file_path, "failed", str(e))
            raise
    
    @check_permission(PermissionLevel.MEDIUM)
    def convert_document(self, file_path: str, output_format: str) -> Dict[str, Any]:
        """Convert document to different format"""
        try:
            validated_path = self._validate_file(file_path)
            
            # Validate output format
            allowed_formats = ['txt', 'md', 'json']
            if output_format.lower() not in allowed_formats:
                raise ValueError(f"Output format {output_format} not allowed")
            
            # Generate output path
            output_path = validated_path.with_suffix(f'.{output_format}')
            
            # VoiceOS native document conversion
            conversion_result = {"success": False, "message": ""}
            
            try:
                # Extract text content
                text_content = ""
                
                if validated_path.suffix.lower() == '.txt':
                    with open(validated_path, 'r', encoding='utf-8') as f:
                        text_content = f.read()
                elif validated_path.suffix.lower() == '.pdf':
                    with open(validated_path, 'rb') as f:
                        pdf_reader = PyPDF2.PdfReader(f)
                        for page in pdf_reader.pages:
                            text_content += page.extract_text() + "\n"
                elif validated_path.suffix.lower() == '.docx':
                    doc = docx.Document(str(validated_path))
                    for paragraph in doc.paragraphs:
                        text_content += paragraph.text + "\n"
                
                # Convert to target format
                if output_format.lower() in ['txt', 'md']:
                    with open(output_path, 'w', encoding='utf-8') as f:
                        f.write(text_content)
                    conversion_result = {
                        "success": True,
                        "message": f"Successfully converted to {output_format}"
                    }
                elif output_format.lower() == 'json':
                    import json
                    json_data = {
                        "source_file": str(validated_path),
                        "content": text_content,
                        "metadata": {
                            "file_type": validated_path.suffix,
                            "conversion_format": output_format
                        }
                    }
                    with open(output_path, 'w', encoding='utf-8') as f:
                        json.dump(json_data, f, indent=2)
                    conversion_result = {
                        "success": True,
                        "message": f"Successfully converted to {output_format}"
                    }
                else:
                    conversion_result = {
                        "success": False,
                        "message": f"Unsupported output format: {output_format}"
                    }
                    
            except Exception as e:
                conversion_result = {
                    "success": False,
                    "message": f"Conversion failed: {str(e)}"
                }
            
            result = {
                "success": conversion_result.get("success", False),
                "output_path": str(output_path),
                "original_format": validated_path.suffix,
                "output_format": output_format,
                "message": conversion_result.get("message", "")
            }
            
            self._log_operation("convert_document", file_path, "success")
            return result
            
        except Exception as e:
            self._log_operation("convert_document", file_path, "failed", str(e))
            raise


# Global instance for tool registry
document_processor = DocumentProcessor()
